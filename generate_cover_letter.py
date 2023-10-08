import argparse
from docx import Document

def generate_cover_letter(company_name, your_name, address, city_country, phone_number, email, github, linkedin):
    doc = Document()
    
    # Function to add a paragraph with bold title and description
    def add_section(title, description):
        p = doc.add_paragraph()
        p.add_run(title).bold = True
        p.add_run("\n" + description)
    
    # Add a title to the document
    doc.add_heading('Cover Letter', level=1)
    
    # Add the intro paragraph
    intro = f"Dear Hiring Managers at {company_name},\n\n"
    doc.add_paragraph(intro)
    
    # Add your information
    your_info = f"Full Name: {your_name}\nAddress: {address}\nCity and Country: {city_country}\n" \
                f"Number: {phone_number}\nEmail: {email}\nGithub: {github}\nLinkedIn: {linkedin}\n\n"
    doc.add_paragraph(your_info)
    
    # Add the main content
    about_me = "About Me:\nDuring my academic journey, I have consistently demonstrated my commitment to learning and growth. " \
               "I firmly believe that sharing knowledge and collaborating with others are the cornerstones of " \
               "personal and professional development. This philosophy has guided me through various experiences, " \
               "including my tenure at ENSAM Casablanca Engineering School and my current pursuit of a Software " \
               "Engineering degree at ENIB.\n\n"
    doc.add_paragraph(about_me)
    
    professional_experience = "Professional Experience:\n"
    doc.add_paragraph(professional_experience)
    
    add_section("Kezakoo (Part-time)", "I have gained valuable experience as a Project Manager and Full Stack " \
                              "Developer at Kezakoo.com, where I accelerated feature delivery by 40% through streamlined " \
                              "deployments and increased student engagement by 30% with a gamification system. Additionally, " \
                              "I improved resource accessibility, reducing search time by 50%, and efficiently managed the " \
                              "development team, achieving a 20% productivity boost. Please refer to my resume for more details " \
                              "on this experience.")

    # Add a brief description of your Miratti experience
    add_section("Miratti Bags (Part-Time)", "Led the next generation of IoT-connected Luxurious Leather bags in Morocco. " \
                "Please refer to my resume for more details on this experience.")
    
    # Add a brief description of your CBI internship
    add_section("CBI (Part-Time Internship)", "Designed and created a groundbreaking automation device using Raspberry Pi " \
                "Zero W and a 4-inch LCD screen to streamline operations for Cisco devices. Implemented versatile connectivity " \
                "options including Serial, Telnet, and SSH, enabling users to collect and manage device logs, execute upgrade procedures, " \
                "and run automated network commands, with logs tracked and stored for reference. Please refer to my resume for more details " \
                "on this experience.")
    
    academic_pursuits = "Academic Pursuits:\nCurrently, I am pursuing a Software Engineer's degree at Brest National School of " \
                        "Engineering, where I continue to expand my skills in big data, frameworks, and software development. " \
                        "I believe in the importance of lifelong learning, and I am committed to staying at the forefront " \
                        "of technological advancements.\n\n"
    doc.add_paragraph(academic_pursuits)
    
    current_situation = "My Current Situation:\nAs I embark on my first year in Brest, I am seeking a part-time job to support " \
                        "myself financially and contribute to my education. I am confident that my technical skills, adaptability, " \
                        "and passion for software engineering make me a valuable asset to any team.\n\n"
    doc.add_paragraph(current_situation)
    
    # Add the closing paragraph
    closing = "I am excited about the opportunity to bring my knowledge and dedication to {company_name} and make meaningful " \
              "contributions to your projects. I am confident that my background in software engineering and my commitment to " \
              "excellence align with your company's mission and values.\n\n" \
              "Thank you for considering my application. I look forward to the possibility of discussing how my skills and " \
              "experiences can contribute to the success of {company_name}.\n\n" \
              "Sincerely,\n\n" \
              "{your_name}"
    doc.add_paragraph(closing.format(company_name=company_name, your_name=your_name))
    
    # Save the document to a file
    doc.save(f"{company_name}_Cover_Letter.docx")

    print(f"Cover letter for {company_name} generated as {company_name}_Cover_Letter.docx.")

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Generate a cover letter.")
    parser.add_argument("--company", type=str, help="Company name")
    args = parser.parse_args()
    
    company_name = args.company
    your_name = "Ayoub Achak"
    address = "2 & 4 rue des Archives 29200 Brest"
    city_country = "Brest, France"
    phone_number = "+33 77 145 2463"
    email = "ayoub.achak01@gmail.com"
    github = "www.github.com/ayoubachak"
    linkedin = "www.linkedin.com/in/ayoubachak"
    
    generate_cover_letter(company_name, your_name, address, city_country, phone_number, email, github, linkedin)

